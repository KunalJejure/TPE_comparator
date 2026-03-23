import io
from typing import List, Dict, Any
from fastapi import HTTPException

def _run_to_html(run) -> str:
    """Convert a single python-docx Run to an HTML fragment preserving formatting."""
    import html as html_mod
    text = html_mod.escape(run.text)
    if not text:
        return ""

    # Apply formatting wrappers
    if run.bold:
        text = f"<strong>{text}</strong>"
    if run.italic:
        text = f"<em>{text}</em>"
    if run.underline:
        text = f"<u>{text}</u>"
    if run.font and run.font.strike:
        text = f"<s>{text}</s>"
    if run.font and run.font.highlight_color:
        text = f'<mark>{text}</mark>'

    return text

def extract_docx_rich(file_bytes: bytes) -> Dict[str, Any]:
    """Extract rich content from a .docx file preserving formatting.
    
    Returns:
        {
            "paragraphs": [...],   # structured paragraph data for matching
            "rich_html": "..."     # full HTML preserving Word formatting
        }
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx is not installed. Run: pip install python-docx",
        )

    doc = Document(io.BytesIO(file_bytes))
    paragraphs: List[Dict[str, Any]] = []
    html_parts: List[str] = []

    # Track list state for bullet/number rendering
    list_open = False
    list_type = ""  # "ul" or "ol"

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style_name = para.style.name if para.style else ""

        # Heading detection
        heading_level = 0
        if style_name.startswith("Heading"):
            try:
                heading_level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                heading_level = 0

        # List detection
        is_list_item = False
        current_list_type = ""
        numPr = para._element.find(qn('w:pPr'))
        if numPr is not None:
            numId_el = numPr.find(qn('w:numPr'))
            if numId_el is not None:
                is_list_item = True
                # Check if numbered or bulleted via abstract num
                current_list_type = "ul"
                ilvl = numId_el.find(qn('w:ilvl'))
                numId = numId_el.find(qn('w:numId'))
                if numId is not None:
                    num_val = numId.get(qn('w:val'))
                    if num_val and int(num_val) > 0:
                        # Heuristic: check style name for numbered
                        if "number" in style_name.lower() or "list number" in style_name.lower():
                            current_list_type = "ol"
        # Also detect by style name
        if "list bullet" in style_name.lower() or "bullet" in style_name.lower():
            is_list_item = True
            current_list_type = "ul"
        elif "list number" in style_name.lower() or "numbered" in style_name.lower():
            is_list_item = True
            current_list_type = "ol"

        # Build rich HTML for this paragraph's runs
        runs_html = ""
        for run in para.runs:
            runs_html += _run_to_html(run)

        if not runs_html and not text:
            # Close any open list before an empty paragraph
            if list_open:
                html_parts.append(f"</{list_type}>")
                list_open = False
            continue

        # Use plain text fallback if runs_html is empty but text exists
        if not runs_html and text:
            import html as html_mod
            runs_html = html_mod.escape(text)

        # Handle list transitions
        if is_list_item:
            if not list_open:
                lt = current_list_type or "ul"
                html_parts.append(f"<{lt} class='sv-doc-list'>")
                list_open = True
                list_type = lt
            elif current_list_type and current_list_type != list_type:
                html_parts.append(f"</{list_type}>")
                html_parts.append(f"<{current_list_type} class='sv-doc-list'>")
                list_type = current_list_type

            html_parts.append(f"<li class='sv-doc-li' data-para-idx='{idx}'>{runs_html}</li>")
        else:
            # Close any open list
            if list_open:
                html_parts.append(f"</{list_type}>")
                list_open = False

            if heading_level:
                tag = f"h{heading_level}"
                html_parts.append(
                    f"<{tag} class='sv-doc-para sv-doc-heading' data-para-idx='{idx}'>"
                    f"{runs_html}</{tag}>"
                )
            else:
                html_parts.append(
                    f"<p class='sv-doc-para' data-para-idx='{idx}'>{runs_html}</p>"
                )

        # Store paragraph data for matching (plain text)
        if text:
            paragraphs.append({
                "index": idx,
                "text": text,
                "style": style_name,
                "heading_level": heading_level,
                "is_list_item": is_list_item,
            })

    # Close any trailing open list
    if list_open:
        html_parts.append(f"</{list_type}>")

    # ── Tables ──
    for table_idx, table in enumerate(doc.tables):
        table_html = "<table class='sv-doc-table'>"
        for row_idx, row in enumerate(table.rows):
            tag = "th" if row_idx == 0 else "td"
            table_html += "<tr>"
            for cell in row.cells:
                cell_text = cell.text.strip()
                import html as html_mod
                table_html += f"<{tag} class='sv-doc-cell'>{html_mod.escape(cell_text)}</{tag}>"
                # Also add table cell text as paragraphs for matching
                if cell_text:
                    para_idx = len(paragraphs) + 10000 + table_idx * 1000 + row_idx
                    paragraphs.append({
                        "index": para_idx,
                        "text": cell_text,
                        "style": "TableCell",
                        "heading_level": 0,
                        "is_list_item": False,
                    })
            table_html += "</tr>"
        table_html += "</table>"
        html_parts.append(table_html)

    return {
        "paragraphs": paragraphs,
        "rich_html": "\n".join(html_parts),
    }
