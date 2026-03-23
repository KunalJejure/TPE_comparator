import os
import docx
from docx2pdf import convert

def test_conversion():
    # 1. Create dummy docx
    doc = docx.Document()
    doc.add_heading('Hello World Word Document', 0)
    doc.add_paragraph('This is a test of python-docx and docx2pdf conversion in the backend.')
    docx_path = "test_doc.docx"
    doc.save(docx_path)
    
    # 2. Convert to PDF
    pdf_path = "test_doc.pdf"
    try:
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            print("SUCCESS: PDF was generated.")
        else:
            print("FAILED: PDF not found.")
    except Exception as e:
        print(f"FAILED with error: {e}")
    finally:
        if os.path.exists(docx_path): os.remove(docx_path)
        if os.path.exists(pdf_path): os.remove(pdf_path)

if __name__ == "__main__":
    test_conversion()
